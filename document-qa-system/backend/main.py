from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import tempfile
from typing import List, Dict, Any, Optional
import uuid
import asyncio
from concurrent.futures import ThreadPoolExecutor
import logging
import json
from datetime import datetime
import re

# Document processing
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
import numpy as np
import pandas as pd
from PIL import Image
import pytesseract
import camelot  # For table extraction from PDFs
import tabula   # Alternative table extraction

# Vector database
from pinecone import Pinecone, ServerlessSpec
import os
from dotenv import load_dotenv

# Sparse retrieval
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load environment variables
load_dotenv()

# LLM
from transformers import AutoTokenizer, AutoModelForCausalLM, BlipProcessor, BlipForConditionalGeneration
import torch

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Multi-Agent Document Q&A API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables for models
embedding_model = None
llm_model = None
llm_tokenizer = None
vision_model = None
vision_processor = None
pc_client = None
executor = ThreadPoolExecutor(max_workers=4)

# In-memory storage for agent metadata and chat sessions
agent_metadata = {}  # {agent_name: {system_prompt: str, created_at: str, ...}}
chat_sessions = {}   # {session_id: {agent_name: str, messages: List, memory: Dict, ...}}
tfidf_vectorizers = {}  # {agent_name: TfidfVectorizer} for sparse retrieval

# Configuration
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
INDEX_NAME = "document-agents"
EMBEDDING_DIM = 384

class QueryRequest(BaseModel):
    question: str
    agent_name: str
    session_id: Optional[str] = None
    memory_enabled: bool = True
    chat_history: List[Dict[str, Any]] = []

class QueryResponse(BaseModel):
    answer: str
    sources: List[str]
    agent_name: str
    session_id: Optional[str] = None

class SystemPromptUpdate(BaseModel):
    system_prompt: str

class ChatMessage(BaseModel):
    type: str  # 'user' or 'ai'
    content: str
    timestamp: str
    sources: Optional[List[str]] = []

@app.on_event("startup")
async def startup_event():
    global embedding_model, llm_model, llm_tokenizer, vision_model, vision_processor, pc_client
    
    logger.info("Loading embedding model...")
    embedding_model = SentenceTransformer(
        'BAAI/bge-small-en-v1.5',
        device="cuda" if torch.cuda.is_available() else "cpu"
    )

    logger.info("Loading vision model for image processing...")
    vision_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    vision_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
    
    logger.info("Loading LLM model...")
    device = "cuda" if torch.cuda.is_available() else "cpu"
    llm_tokenizer = AutoTokenizer.from_pretrained("deepseek-ai/deepseek-coder-1.3b-instruct")
    llm_model = AutoModelForCausalLM.from_pretrained(
        "deepseek-ai/deepseek-coder-1.3b-instruct",
        torch_dtype=torch.float16 if device == "cuda" else torch.float32
    ).to(device)

    logger.info("Initializing Pinecone...")
    pc_client = Pinecone(api_key=PINECONE_API_KEY)
    
    # Create index if it doesn't exist
    try:
        index_info = pc_client.describe_index(INDEX_NAME)
        logger.info(f"Index {INDEX_NAME} already exists")
    except:
        logger.info(f"Creating index {INDEX_NAME}")
        pc_client.create_index(
            name=INDEX_NAME,
            dimension=EMBEDDING_DIM,
            metric="cosine",
            spec=ServerlessSpec(
                cloud="aws",
                region="us-east-1"
            )
        )

def extract_tables_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extract tables from PDF using camelot"""
    tables = []
    try:
        # Use camelot for table extraction
        camelot_tables = camelot.read_pdf(file_path, pages='all')
        for i, table in enumerate(camelot_tables):
            if table.df is not None and not table.df.empty:
                # Convert table to text representation
                table_text = table.df.to_string(index=False)
                tables.append({
                    'type': 'table',
                    'content': table_text,
                    'page': i + 1,
                    'table_index': i
                })
    except Exception as e:
        logger.warning(f"Camelot table extraction failed: {e}")
        
        # Fallback to tabula
        try:
            tabula_tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
            for i, table in enumerate(tabula_tables):
                if table is not None and not table.empty:
                    table_text = table.to_string(index=False)
                    tables.append({
                        'type': 'table',
                        'content': table_text,
                        'page': i + 1,
                        'table_index': i
                    })
        except Exception as e2:
            logger.warning(f"Tabula table extraction also failed: {e2}")
    
    return tables

def extract_images_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extract and process images from PDF"""
    images = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    # Convert to PIL Image
                    img_data = pix.tobytes("ppm")
                    pil_image = Image.open(io.BytesIO(img_data))
                    
                    # Generate caption using vision model
                    inputs = vision_processor(pil_image, return_tensors="pt")
                    out = vision_model.generate(**inputs, max_length=50)
                    caption = vision_processor.decode(out[0], skip_special_tokens=True)
                    
                    # OCR text extraction
                    ocr_text = pytesseract.image_to_string(pil_image)
                    
                    images.append({
                        'type': 'image',
                        'content': f"Image caption: {caption}\nExtracted text: {ocr_text}",
                        'page': page_num + 1,
                        'image_index': img_index
                    })
                
                pix = None
        
        doc.close()
    except Exception as e:
        logger.warning(f"Image extraction failed: {e}")
    
    return images

def extract_text_from_pdf(file_path: str) -> tuple[str, List[Dict], List[Dict]]:
    """Extract text, tables, and images from PDF file"""
    text = ""
    tables = []
    images = []
    
    # Extract regular text
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    
    # Extract tables
    tables = extract_tables_from_pdf(file_path)
    
    # Extract images
    images = extract_images_from_pdf(file_path)
    
    return text, tables, images

def extract_text_from_docx(file_path: str) -> tuple[str, List[Dict], List[Dict]]:
    """Extract text and tables from DOCX file"""
    doc = docx.Document(file_path)
    text = ""
    tables = []
    
    # Extract paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            # Convert to DataFrame for better formatting
            df = pd.DataFrame(table_data[1:], columns=table_data[0])
            table_text = df.to_string(index=False)
            tables.append({
                'type': 'table',
                'content': table_text,
                'table_index': i
            })
    
    return text, tables, []  # No image extraction for DOCX yet

def extract_text_from_txt(file_path: str) -> tuple[str, List[Dict], List[Dict]]:
    """Extract text from TXT file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read(), [], []

def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
    """Split text into overlapping chunks with metadata"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk_text = ' '.join(words[i:i + chunk_size])
        if chunk_text.strip():
            chunks.append({
                'type': 'text',
                'content': chunk_text,
                'chunk_index': len(chunks),
                'word_start': i,
                'word_end': min(i + chunk_size, len(words))
            })
        
        if i + chunk_size >= len(words):
            break
    
    return chunks

def prepare_all_chunks(text: str, tables: List[Dict], images: List[Dict]) -> List[Dict[str, Any]]:
    """Combine text chunks, tables, and images into unified chunk list"""
    all_chunks = []
    
    # Add text chunks
    text_chunks = split_text_into_chunks(text)
    all_chunks.extend(text_chunks)
    
    # Add table chunks
    for table in tables:
        all_chunks.append({
            'type': 'table',
            'content': f"TABLE DATA:\n{table['content']}",
            'chunk_index': len(all_chunks),
            'page': table.get('page'),
            'table_index': table.get('table_index')
        })
    
    # Add image chunks
    for image in images:
        all_chunks.append({
            'type': 'image',
            'content': f"IMAGE CONTENT:\n{image['content']}",
            'chunk_index': len(all_chunks),
            'page': image.get('page'),
            'image_index': image.get('image_index')
        })
    
    return all_chunks

def generate_embeddings(texts: List[str]) -> np.ndarray:
    """Generate embeddings for text chunks"""
    return embedding_model.encode(texts, normalize_embeddings=True)

def create_sparse_index(chunks: List[Dict[str, Any]], agent_name: str):
    """Create TF-IDF sparse index for agent"""
    global tfidf_vectorizers
    
    chunk_contents = [chunk['content'] for chunk in chunks]
    vectorizer = TfidfVectorizer(max_features=5000, stop_words='english')
    vectorizer.fit(chunk_contents)
    tfidf_vectorizers[agent_name] = vectorizer

def hybrid_search(query: str, agent_name: str, top_k: int = 5) -> List[Dict]:
    """Perform hybrid sparse + dense search"""
    # Dense search (existing)
    question_embedding = embedding_model.encode([query], normalize_embeddings=True)[0]
    
    index = pc_client.Index(INDEX_NAME)
    dense_results = index.query(
        vector=question_embedding.tolist(),
        top_k=top_k * 2,  # Get more candidates
        include_metadata=True,
        namespace=agent_name
    )
    
    # Sparse search
    sparse_scores = {}
    if agent_name in tfidf_vectorizers:
        query_tfidf = tfidf_vectorizers[agent_name].transform([query])
        
        # Get all chunks for this agent for sparse comparison
        # This is simplified - in production, you'd want a more efficient approach
        all_results = index.query(
            vector=[0] * EMBEDDING_DIM,
            top_k=10000,  # Large number to get all
            include_metadata=True,
            namespace=agent_name
        )
        
        chunk_contents = [match.metadata['text'] for match in all_results.matches]
        if chunk_contents:
            chunk_tfidf = tfidf_vectorizers[agent_name].transform(chunk_contents)
            sparse_similarities = cosine_similarity(query_tfidf, chunk_tfidf).flatten()
            
            for i, match in enumerate(all_results.matches):
                sparse_scores[match.id] = sparse_similarities[i]
    
    # Combine scores (weighted)
    dense_weight = 0.7
    sparse_weight = 0.3
    
    combined_results = []
    for match in dense_results.matches:
        dense_score = match.score
        sparse_score = sparse_scores.get(match.id, 0)
        combined_score = dense_weight * dense_score + sparse_weight * sparse_score
        
        combined_results.append({
            'id': match.id,
            'score': combined_score,
            'metadata': match.metadata
        })
    
    # Sort by combined score and return top_k
    combined_results.sort(key=lambda x: x['score'], reverse=True)
    return combined_results[:top_k]

def get_chat_memory_context(session_id: str, memory_enabled: bool) -> str:
    """Get relevant context from chat memory"""
    if not memory_enabled or session_id not in chat_sessions:
        return ""
    
    session = chat_sessions[session_id]
    recent_messages = session.get('messages', [])[-10:]  # Last 10 messages
    
    memory_context = []
    for msg in recent_messages:
        if msg['type'] == 'user':
            memory_context.append(f"User: {msg['content']}")
        else:
            memory_context.append(f"Assistant: {msg['content']}")
    
    return "\n".join(memory_context) if memory_context else ""

def update_chat_memory(session_id: str, user_message: str, ai_response: str, agent_name: str):
    """Update chat memory for session"""
    if session_id not in chat_sessions:
        chat_sessions[session_id] = {
            'agent_name': agent_name,
            'messages': [],
            'created_at': datetime.now().isoformat(),
            'last_activity': datetime.now().isoformat()
        }
    
    session = chat_sessions[session_id]
    session['messages'].extend([
        {
            'type': 'user',
            'content': user_message,
            'timestamp': datetime.now().isoformat()
        },
        {
            'type': 'ai',
            'content': ai_response,
            'timestamp': datetime.now().isoformat()
        }
    ])
    session['last_activity'] = datetime.now().isoformat()

def generate_answer(question: str, context: str, system_prompt: str, memory_context: str = "") -> str:
    """Generate answer using LLM with system prompt and memory"""
    
    # Build comprehensive prompt
    prompt_parts = []
    
    if system_prompt:
        prompt_parts.append(f"System: {system_prompt}")
    
    if memory_context:
        prompt_parts.append(f"Previous conversation:\n{memory_context}")
    
    prompt_parts.extend([
        f"Context from documents:\n{context}",
        f"Current question: {question}",
        "Answer:"
    ])
    
    prompt = "\n\n".join(prompt_parts)
    
    inputs = llm_tokenizer(prompt, return_tensors="pt", max_length=2048, truncation=True)
    
    device = "cuda" if torch.cuda.is_available() else "cpu"
    if torch.cuda.is_available():
        inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        outputs = llm_model.generate(
            **inputs,
            max_new_tokens=300,
            temperature=0.7,
            do_sample=True,
            pad_token_id=llm_tokenizer.eos_token_id
        )
    
    response = llm_tokenizer.decode(outputs[0], skip_special_tokens=True)
    # Extract only the answer part
    answer_start = response.find("Answer:") + len("Answer:")
    answer = response[answer_start:].strip()
    
    return answer

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    agent_name: str = Form(...),
    system_prompt: str = Form("You are a helpful AI assistant that answers questions based on the provided documents. Be accurate, concise, and cite relevant information when possible.")
):
    """Upload and process a document for a specific agent"""
    
    if not agent_name.strip():
        raise HTTPException(status_code=400, detail="Agent name is required")
    
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.txt'}
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # Extract text, tables, and images based on file type
        if file_extension == '.pdf':
            text, tables, images = extract_text_from_pdf(temp_file_path)
        elif file_extension == '.docx':
            text, tables, images = extract_text_from_docx(temp_file_path)
        elif file_extension == '.txt':
            text, tables, images = extract_text_from_txt(temp_file_path)
        
        # Clean up temp file
        os.unlink(temp_file_path)
        
        if not text.strip() and not tables and not images:
            raise HTTPException(status_code=400, detail="No content found in the document")
        
        # Prepare all chunks (text + tables + images)
        all_chunks = prepare_all_chunks(text, tables, images)
        logger.info(f"Prepared {len(all_chunks)} total chunks ({len([c for c in all_chunks if c['type'] == 'text'])} text, {len([c for c in all_chunks if c['type'] == 'table'])} tables, {len([c for c in all_chunks if c['type'] == 'image'])} images)")
        
        # Generate embeddings for all chunks
        chunk_contents = [chunk['content'] for chunk in all_chunks]
        embeddings = generate_embeddings(chunk_contents)
        
        # Create sparse index for this agent
        create_sparse_index(all_chunks, agent_name)
        
        # Store in Pinecone
        index = pc_client.Index(INDEX_NAME)
        
        vectors_to_upsert = []
        for i, (chunk, embedding) in enumerate(zip(all_chunks, embeddings)):
            vector_id = f"{agent_name}_{file.filename}_{i}_{uuid.uuid4().hex[:8]}"
            vectors_to_upsert.append({
                "id": vector_id,
                "values": embedding.tolist(),
                "metadata": {
                    "text": chunk['content'],
                    "filename": file.filename,
                    "chunk_index": chunk['chunk_index'],
                    "agent_name": agent_name,
                    "content_type": chunk['type'],
                    "page": chunk.get('page'),
                    "table_index": chunk.get('table_index'),
                    "image_index": chunk.get('image_index')
                }
            })
        
        # Upsert vectors to Pinecone with namespace
        index.upsert(
            vectors=vectors_to_upsert,
            namespace=agent_name
        )
        
        # Store agent metadata
        if agent_name not in agent_metadata:
            agent_metadata[agent_name] = {
                'system_prompt': system_prompt,
                'created_at': datetime.now().isoformat(),
                'documents': []
            }
        
        agent_metadata[agent_name]['documents'].append({
            'filename': file.filename,
            'uploaded_at': datetime.now().isoformat(),
            'chunks': len(all_chunks),
            'tables': len(tables),
            'images': len(images)
        })
        
        logger.info(f"Successfully uploaded {file.filename} for agent {agent_name}")
        
        return {
            "message": f"Successfully uploaded {file.filename}",
            "agent_name": agent_name,
            "chunks_processed": len(all_chunks),
            "filename": file.filename,
            "tables_extracted": len(tables),
            "images_processed": len(images)
        }
        
    except Exception as e:
        logger.error(f"Error processing upload: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post("/query", response_model=QueryResponse)
async def query_agent(request: QueryRequest):
    """Query a specific agent with a question"""
    
    if request.agent_name not in agent_metadata:
        raise HTTPException(status_code=404, detail=f"Agent '{request.agent_name}' not found")
    
    try:
        # Get memory context if enabled
        memory_context = ""
        if request.memory_enabled and request.session_id:
            memory_context = get_chat_memory_context(request.session_id, request.memory_enabled)
        
        # Perform hybrid search
        search_results = hybrid_search(request.question, request.agent_name, top_k=5)
        
        if not search_results:
            return QueryResponse(
                answer="I cannot find any relevant information in the documents for this agent.",
                sources=[],
                agent_name=request.agent_name,
                session_id=request.session_id
            )
        
        # Prepare context from retrieved chunks
        context_chunks = []
        sources = set()
        
        for result in search_results:
            if result['score'] > 0.6:  # Lower threshold for hybrid search
                context_chunks.append(result['metadata']['text'])
                sources.add(result['metadata']['filename'])
                
                # Add content type info for better context
                content_type = result['metadata'].get('content_type', 'text')
                if content_type != 'text':
                    context_chunks[-1] = f"[{content_type.upper()}] {context_chunks[-1]}"

        # Fallback: take best match if nothing passed threshold
        if not context_chunks and search_results:
            best_result = search_results[0]
            context_chunks.append(best_result['metadata']['text'])
            sources.add(best_result['metadata']['filename'])
        
        if not context_chunks:
            return QueryResponse(
                answer="I cannot find sufficiently relevant information to answer your question.",
                sources=[],
                agent_name=request.agent_name,
                session_id=request.session_id
            )
        
        context = "\n\n".join(context_chunks)
        system_prompt = agent_metadata[request.agent_name]['system_prompt']
        
        # Generate answer using LLM with system prompt and memory
        answer = await asyncio.get_event_loop().run_in_executor(
            executor, generate_answer, request.question, context, system_prompt, memory_context
        )
        
        # Update chat memory if session provided
        if request.session_id and request.memory_enabled:
            update_chat_memory(request.session_id, request.question, answer, request.agent_name)
        
        return QueryResponse(
            answer=answer,
            sources=list(sources),
            agent_name=request.agent_name,
            session_id=request.session_id
        )
        
    except Exception as e:
        logger.error(f"Error processing query: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

@app.get("/agents")
async def list_agents():
    """List all available agents with metadata"""
    try:
        index = pc_client.Index(INDEX_NAME)
        
        # Get index stats to see all namespaces
        stats = index.describe_index_stats()
        namespaces = list(stats.namespaces.keys()) if stats.namespaces else []
        
        agents = []
        for namespace in namespaces:
            if namespace:  # Skip empty namespace
                namespace_stats = stats.namespaces[namespace]
                agent_info = {
                    "name": namespace,
                    "document_count": namespace_stats.vector_count,
                    "system_prompt": agent_metadata.get(namespace, {}).get('system_prompt', ''),
                    "created_at": agent_metadata.get(namespace, {}).get('created_at', ''),
                    "documents": agent_metadata.get(namespace, {}).get('documents', [])
                }
                agents.append(agent_info)
        
        return {"agents": agents}
        
    except Exception as e:
        logger.error(f"Error listing agents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing agents: {str(e)}")

@app.put("/agents/{agent_name}/system-prompt")
async def update_agent_system_prompt(agent_name: str, update: SystemPromptUpdate):
    """Update system prompt for an agent"""
    if agent_name not in agent_metadata:
        raise HTTPException(status_code=404, detail=f"Agent '{agent_name}' not found")
    
    try:
        agent_metadata[agent_name]['system_prompt'] = update.system_prompt
        agent_metadata[agent_name]['updated_at'] = datetime.now().isoformat()
        
        return {
            "message": f"Successfully updated system prompt for agent '{agent_name}'",
            "system_prompt": update.system_prompt
        }
        
    except Exception as e:
        logger.error(f"Error updating system prompt: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error updating system prompt: {str(e)}")

@app.get("/agents/{agent_name}")
async def get_agent_details(agent_name: str):
    """Get detailed information about a specific agent"""
    if agent_name not in agent_metadata:
        raise HTTPException(status_code=404, detail=f"Agent '{agent_name}' not found")
    
    try:
        # Get vector count from Pinecone
        index = pc_client.Index(INDEX_NAME)
        stats = index.describe_index_stats()
        vector_count = stats.namespaces.get(agent_name, {}).get('vector_count', 0)
        
        agent_info = agent_metadata[agent_name].copy()
        agent_info['document_count'] = vector_count
        agent_info['name'] = agent_name
        
        return agent_info
        
    except Exception as e:
        logger.error(f"Error getting agent details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting agent details: {str(e)}")

@app.delete("/agents/{agent_name}")
async def delete_agent(agent_name: str):
    """Delete an agent and all its documents"""
    try:
        index = pc_client.Index(INDEX_NAME)
        
        # Delete all vectors in the namespace
        index.delete(delete_all=True, namespace=agent_name)
        
        # Remove from metadata
        if agent_name in agent_metadata:
            del agent_metadata[agent_name]
        
        # Remove sparse index
        if agent_name in tfidf_vectorizers:
            del tfidf_vectorizers[agent_name]
        
        # Remove related chat sessions
        sessions_to_remove = [sid for sid, session in chat_sessions.items() 
                             if session.get('agent_name') == agent_name]
        for sid in sessions_to_remove:
            del chat_sessions[sid]
        
        return {"message": f"Successfully deleted agent '{agent_name}' and all its documents"}
        
    except Exception as e:
        logger.error(f"Error deleting agent: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting agent: {str(e)}")

@app.get("/sessions")
async def list_chat_sessions():
    """List all chat sessions"""
    sessions = []
    for session_id, session_data in chat_sessions.items():
        session_info = {
            'id': session_id,
            'agent_name': session_data['agent_name'],
            'created_at': session_data['created_at'],
            'last_activity': session_data['last_activity'],
            'message_count': len(session_data.get('messages', []))
        }
        sessions.append(session_info)
    
    return {"sessions": sessions}

@app.get("/sessions/{session_id}")
async def get_chat_session(session_id: str):
    """Get specific chat session details"""
    if session_id not in chat_sessions:
        raise HTTPException(status_code=404, detail=f"Session '{session_id}' not found")
    
    return chat_sessions[session_id]

@app.delete("/sessions/{session_id}")
async def delete_chat_session(session_id: str):
    """Delete a specific chat session"""
    if session_id not in chat_sessions:
        raise HTTPException(status_code=404, detail=f"Session '{session_id}' not found")
    
    del chat_sessions[session_id]
    return {"message": f"Successfully deleted session '{session_id}'"}

@app.post("/sessions")
async def create_chat_session(agent_name: str):
    """Create a new chat session for an agent"""
    if agent_name not in agent_metadata:
        raise HTTPException(status_code=404, detail=f"Agent '{agent_name}' not found")
    
    session_id = f"session_{uuid.uuid4().hex[:8]}"
    chat_sessions[session_id] = {
        'agent_name': agent_name,
        'messages': [],
        'created_at': datetime.now().isoformat(),
        'last_activity': datetime.now().isoformat()
    }
    
    return {
        "session_id": session_id,
        "agent_name": agent_name,
        "message": f"Created new session for agent '{agent_name}'"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "embedding_model": "BAAI/bge-small-en-v1.5",
        "llm_model": "deepseek-ai/deepseek-coder-1.3b-instruct",
        "vision_model": "Salesforce/blip-image-captioning-base",
        "vector_db": "Pinecone",
        "features": {
            "hybrid_search": True,
            "table_extraction": True,
            "image_processing": True,
            "persistent_memory": True,
            "custom_system_prompts": True
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)